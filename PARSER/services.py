import os
import re
import json
import hashlib
from typing import Dict, List, Tuple, Any
from pathlib import Path
from django.conf import settings
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile

# Importaciones para parseo de documentos
try:
    from docx import Document as DocxDocument
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    DocxDocument = None

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    from openpyxl.utils import get_column_letter
except ImportError:
    openpyxl = None

try:
    import PyPDF2
    from PyPDF2 import PdfReader
except ImportError:
    PyPDF2 = None
    PdfReader = None

try:
    from PIL import Image
except ImportError:
    Image = None

class DocumentParserService:
    """
    Servicio principal para parsear documentos y convertirlos a HTML/CSS/JS
    """
    
    def __init__(self):
        self.supported_formats = ['.docx', '.xlsx', '.xls', '.pdf']
        self.placeholder_pattern = r'\{\{([^}]+)\}\}'
        
    def parse_document(self, file_path: str, document_type: str) -> Dict[str, Any]:
        """
        Parsea un documento y retorna el contenido estructurado
        """
        try:
            if document_type == 'docx':
                return self._parse_docx(file_path)
            elif document_type in ['xlsx', 'xls']:
                return self._parse_excel(file_path)
            elif document_type == 'pdf':
                return self._parse_pdf(file_path)
            else:
                raise ValueError(f"Formato no soportado: {document_type}")
        except Exception as e:
            return {
                'success': False,
                'error': str(e),
                'content': None
            }
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """
        Parsea un documento DOCX y extrae contenido, estilos y placeholders
        """
        if not DocxDocument:
            raise ImportError("python-docx no está instalado")
        
        doc = DocxDocument(file_path)
        
        # Extraer contenido y estilos
        content_data = {
            'paragraphs': [],
            'tables': [],
            'images': [],
            'styles': {},
            'placeholders': []
        }
        
        # Procesar párrafos
        for para in doc.paragraphs:
            para_data = self._extract_paragraph_data(para)
            if para_data['text'].strip():
                content_data['paragraphs'].append(para_data)
        
        # Procesar tablas
        for table in doc.tables:
            table_data = self._extract_table_data(table)
            content_data['tables'].append(table_data)
        
        # Generar HTML/CSS/JS
        html_content = self._generate_html_from_docx(content_data)
        css_content = self._generate_css_from_docx(content_data)
        js_content = self._generate_js_template()
        
        # Detectar placeholders
        placeholders = self._detect_placeholders(html_content)
        
        return {
            'success': True,
            'content': {
                'raw_data': content_data,
                'html': html_content,
                'css': css_content,
                'js': js_content,
                'placeholders': placeholders
            }
        }
    
    def _extract_paragraph_data(self, paragraph) -> Dict[str, Any]:
        """
        Extrae datos de un párrafo incluyendo formato y estilos
        """
        para_data = {
            'text': paragraph.text,
            'style': paragraph.style.name if paragraph.style else 'Normal',
            'alignment': self._get_alignment(paragraph.alignment),
            'runs': []
        }
        
        # Procesar runs (fragmentos de texto con formato)
        for run in paragraph.runs:
            run_data = {
                'text': run.text,
                'bold': run.bold,
                'italic': run.italic,
                'underline': run.underline,
                'font_name': run.font.name,
                'font_size': run.font.size.pt if run.font.size else None,
                'color': self._get_color(run.font.color),
                'is_placeholder': self._is_red_text(run.font.color)
            }
            para_data['runs'].append(run_data)
        
        return para_data
    
    def _extract_table_data(self, table) -> Dict[str, Any]:
        """
        Extrae datos de una tabla incluyendo celdas y estilos
        """
        table_data = {
            'rows': [],
            'style': table.style.name if table.style else None
        }
        
        for row in table.rows:
            row_data = {'cells': []}
            for cell in row.cells:
                cell_data = {
                    'text': cell.text,
                    'paragraphs': []
                }
                
                # Procesar párrafos dentro de la celda
                for para in cell.paragraphs:
                    cell_data['paragraphs'].append(self._extract_paragraph_data(para))
                
                row_data['cells'].append(cell_data)
            table_data['rows'].append(row_data)
        
        return table_data
    
    def _parse_excel(self, file_path: str) -> Dict[str, Any]:
        """
        Parsea un archivo Excel y extrae contenido, estilos y placeholders
        """
        if not openpyxl:
            raise ImportError("openpyxl no está instalado")
        
        workbook = openpyxl.load_workbook(file_path, data_only=False)
        
        content_data = {
            'sheets': [],
            'styles': {},
            'placeholders': []
        }
        
        # Procesar cada hoja
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_data = self._extract_sheet_data(sheet)
            content_data['sheets'].append(sheet_data)
        
        # Generar HTML/CSS/JS
        html_content = self._generate_html_from_excel(content_data)
        css_content = self._generate_css_from_excel(content_data)
        js_content = self._generate_js_template()
        
        # Detectar placeholders
        placeholders = self._detect_placeholders(html_content)
        
        return {
            'success': True,
            'content': {
                'raw_data': content_data,
                'html': html_content,
                'css': css_content,
                'js': js_content,
                'placeholders': placeholders
            }
        }
    
    def _extract_sheet_data(self, sheet) -> Dict[str, Any]:
        """
        Extrae datos de una hoja de Excel
        """
        sheet_data = {
            'name': sheet.title,
            'cells': {},
            'merged_cells': [],
            'dimensions': {
                'max_row': sheet.max_row,
                'max_column': sheet.max_column
            }
        }
        
        # Extraer celdas con contenido
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    cell_data = {
                        'value': str(cell.value),
                        'coordinate': cell.coordinate,
                        'row': cell.row,
                        'column': cell.column,
                        'font': self._extract_excel_font(cell.font),
                        'fill': self._extract_excel_fill(cell.fill),
                        'border': self._extract_excel_border(cell.border),
                        'alignment': self._extract_excel_alignment(cell.alignment),
                        'is_placeholder': self._is_red_excel_cell(cell)
                    }
                    sheet_data['cells'][cell.coordinate] = cell_data
        
        # Extraer celdas combinadas
        for merged_range in sheet.merged_cells.ranges:
            sheet_data['merged_cells'].append(str(merged_range))
        
        return sheet_data
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """
        Parsea un archivo PDF y extrae texto (funcionalidad básica)
        """
        if not PdfReader:
            raise ImportError("PyPDF2 no está instalado")
        
        with open(file_path, 'rb') as file:
            pdf_reader = PdfReader(file)
            
            content_data = {
                'pages': [],
                'metadata': pdf_reader.metadata,
                'num_pages': len(pdf_reader.pages)
            }
            
            # Extraer texto de cada página
            for page_num, page in enumerate(pdf_reader.pages):
                page_text = page.extract_text()
                content_data['pages'].append({
                    'page_number': page_num + 1,
                    'text': page_text
                })
        
        # Generar HTML básico para PDF
        html_content = self._generate_html_from_pdf(content_data)
        css_content = self._generate_css_from_pdf()
        js_content = self._generate_js_template()
        
        # Detectar placeholders en el texto
        placeholders = self._detect_placeholders_in_text('\n'.join([page['text'] for page in content_data['pages']]))
        
        return {
            'success': True,
            'content': {
                'raw_data': content_data,
                'html': html_content,
                'css': css_content,
                'js': js_content,
                'placeholders': placeholders
            }
        }
    
    def _generate_html_from_docx(self, content_data: Dict) -> str:
        """
        Genera HTML completo con CSS y JS embebidos a partir de datos extraídos de DOCX
        """
        # Generar CSS embebido
        css_content = self._generate_css_from_docx(content_data)
        js_content = self._generate_js_template()
        
        html_parts = ['<!DOCTYPE html>', '<html lang="es">', '<head>']
        html_parts.append('<meta charset="UTF-8">')
        html_parts.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_parts.append('<title>Plantilla de Documento</title>')
        
        # Embeber CSS
        html_parts.append('<style>')
        html_parts.append(css_content)
        html_parts.append('</style>')
        
        html_parts.extend(['</head>', '<body>', '<div class="document-container">'])
        
        # Generar párrafos
        for para in content_data['paragraphs']:
            html_parts.append(self._paragraph_to_html(para))
        
        # Generar tablas
        for table in content_data['tables']:
            html_parts.append(self._table_to_html(table))
        
        html_parts.append('</div>')
        
        # Embeber JavaScript
        html_parts.append('<script>')
        html_parts.append(js_content)
        html_parts.append('</script>')
        
        html_parts.extend(['</body>', '</html>'])
        
        return '\n'.join(html_parts)
    
    def _generate_html_from_excel(self, content_data: Dict) -> str:
        """
        Genera HTML completo con CSS y JS embebidos a partir de datos extraídos de Excel
        """
        # Generar CSS y JS embebidos
        css_content = self._generate_css_from_excel(content_data)
        js_content = self._generate_js_template()
        
        html_parts = ['<!DOCTYPE html>', '<html lang="es">', '<head>']
        html_parts.append('<meta charset="UTF-8">')
        html_parts.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_parts.append('<title>Plantilla de Hoja de Cálculo</title>')
        
        # Embeber CSS
        html_parts.append('<style>')
        html_parts.append(css_content)
        html_parts.append('</style>')
        
        html_parts.extend(['</head>', '<body>'])
        
        # Generar cada hoja como una tabla
        for sheet in content_data['sheets']:
            html_parts.append(f'<div class="sheet-container" data-sheet="{sheet["name"]}">')
            html_parts.append(f'<h2 class="sheet-title">{sheet["name"]}</h2>')
            html_parts.append(self._sheet_to_html_table(sheet))
            html_parts.append('</div>')
        
        # Embeber JavaScript
        html_parts.append('<script>')
        html_parts.append(js_content)
        html_parts.append('</script>')
        
        html_parts.extend(['</body>', '</html>'])
        
        return '\n'.join(html_parts)
    
    def _generate_html_from_pdf(self, content_data: Dict) -> str:
        """
        Genera HTML completo con CSS y JS embebidos a partir de datos extraídos de PDF
        """
        # Generar CSS y JS embebidos
        css_content = self._generate_css_from_pdf()
        js_content = self._generate_js_template()
        
        html_parts = ['<!DOCTYPE html>', '<html lang="es">', '<head>']
        html_parts.append('<meta charset="UTF-8">')
        html_parts.append('<meta name="viewport" content="width=device-width, initial-scale=1.0">')
        html_parts.append('<title>Plantilla de PDF</title>')
        
        # Embeber CSS
        html_parts.append('<style>')
        html_parts.append(css_content)
        html_parts.append('</style>')
        
        html_parts.extend(['</head>', '<body>', '<div class="pdf-container">'])
        
        # Generar cada página
        for page in content_data['pages']:
            html_parts.append(f'<div class="pdf-page" data-page="{page["page_number"]}">')
            # Convertir texto a párrafos y detectar placeholders
            text_with_placeholders = self._convert_text_to_html_with_placeholders(page['text'])
            html_parts.append(text_with_placeholders)
            html_parts.append('</div>')
        
        html_parts.append('</div>')
        
        # Embeber JavaScript
        html_parts.append('<script>')
        html_parts.append(js_content)
        html_parts.append('</script>')
        
        html_parts.extend(['</body>', '</html>'])
        
        return '\n'.join(html_parts)
    
    def _generate_css_from_docx(self, content_data: Dict) -> str:
        """
        Genera CSS a partir de estilos extraídos de DOCX
        """
        css_parts = [
            '/* Estilos generados automáticamente desde DOCX */',
            'body { font-family: "Calibri", Arial, sans-serif; margin: 0; padding: 20px; }',
            '.document-container { max-width: 210mm; margin: 0 auto; background: white; padding: 20mm; box-shadow: 0 0 10px rgba(0,0,0,0.1); }',
            '',
            '/* Estilos de párrafos */',
            '.paragraph { margin-bottom: 12px; line-height: 1.15; }',
            '.paragraph.center { text-align: center; }',
            '.paragraph.right { text-align: right; }',
            '.paragraph.justify { text-align: justify; }',
            '',
            '/* Estilos de texto */',
            '.bold { font-weight: bold; }',
            '.italic { font-style: italic; }',
            '.underline { text-decoration: underline; }',
            '.red-text { color: #ff0000; }',
            '',
            '/* Estilos de tablas */',
            '.document-table { width: 100%; border-collapse: collapse; margin: 12px 0; }',
            '.document-table td, .document-table th { border: 1px solid #000; padding: 8px; vertical-align: top; }',
            '',
            '/* Placeholders */',
            '.placeholder { color: #ff0000; font-weight: normal; }',
            '.placeholder:hover { background-color: #ffe599; }'
        ]
        
        return '\n'.join(css_parts)
    
    def _generate_css_from_excel(self, content_data: Dict) -> str:
        """
        Genera CSS a partir de estilos extraídos de Excel
        """
        css_parts = [
            '/* Estilos generados automáticamente desde Excel */',
            'body { font-family: "Calibri", Arial, sans-serif; margin: 0; padding: 20px; }',
            '.sheet-container { margin-bottom: 30px; }',
            '.sheet-title { font-size: 18px; font-weight: bold; margin-bottom: 10px; color: #333; }',
            '',
            '/* Estilos de tabla Excel */',
            '.excel-table { border-collapse: collapse; width: 100%; }',
            '.excel-table td, .excel-table th { border: 1px solid #d0d7de; padding: 6px 8px; text-align: left; vertical-align: top; }',
            '.excel-table th { background-color: #f6f8fa; font-weight: bold; }',
            '',
            '/* Estilos de celdas */',
            '.cell-bold { font-weight: bold; }',
            '.cell-italic { font-style: italic; }',
            '.cell-center { text-align: center; }',
            '.cell-right { text-align: right; }',
            '',
            '/* Placeholders */',
            '.placeholder { color: #ff0000; font-weight: normal; }',
            '.placeholder:hover { background-color: #ffe599; }'
        ]
        
        return '\n'.join(css_parts)
    
    def _generate_css_from_pdf(self) -> str:
        """
        Genera CSS básico para contenido de PDF
        """
        css_parts = [
            '/* Estilos generados automáticamente desde PDF */',
            'body { font-family: "Times New Roman", serif; margin: 0; padding: 20px; line-height: 1.6; }',
            '.pdf-container { max-width: 210mm; margin: 0 auto; }',
            '.pdf-page { margin-bottom: 30px; padding: 20px; border: 1px solid #ddd; background: white; }',
            '.pdf-page p { margin-bottom: 12px; }',
            '',
            '/* Placeholders */',
            '.placeholder { color: #ff0000; font-weight: normal; }',
            '.placeholder:hover { background-color: #ffe599; }'
        ]
        
        return '\n'.join(css_parts)
    
    def _generate_js_template(self) -> str:
        """
        Genera archivo JavaScript básico para la plantilla
        """
        js_content = '''// Script generado automáticamente para la plantilla
// Este archivo está preparado para manejar la lógica de placeholders

document.addEventListener('DOMContentLoaded', function() {
    console.log('Plantilla cargada correctamente');
    
    // Inicializar placeholders
    initializePlaceholders();
});

function initializePlaceholders() {
    const placeholders = document.querySelectorAll('.placeholder');
    
    placeholders.forEach(placeholder => {
        placeholder.addEventListener('click', function() {
            console.log('Placeholder clickeado:', this.textContent);
        });
    });
}

// Función para rellenar placeholders con datos
function fillPlaceholders(data) {
    Object.keys(data).forEach(key => {
        const elements = document.querySelectorAll(`[data-placeholder="${key}"]`);
        elements.forEach(element => {
            element.textContent = data[key];
        });
    });
}

// Función para obtener todos los placeholders
function getPlaceholders() {
    const placeholders = document.querySelectorAll('.placeholder');
    const placeholderList = [];
    
    placeholders.forEach(placeholder => {
        const key = placeholder.getAttribute('data-placeholder');
        if (key && !placeholderList.includes(key)) {
            placeholderList.push(key);
        }
    });
    
    return placeholderList;
}'''
        
        return js_content
    
    def _detect_placeholders(self, html_content: str) -> List[Dict[str, Any]]:
        """
        Detecta placeholders en el contenido HTML
        """
        placeholders = []
        pattern = r'\{\{([^}]+)\}\}'
        matches = re.finditer(pattern, html_content)
        
        for match in matches:
            placeholder_name = match.group(1).strip()
            placeholders.append({
                'name': placeholder_name,
                'type': self._infer_placeholder_type(placeholder_name),
                'position': match.span(),
                'original_text': match.group(0)
            })
        
        return placeholders
    
    def _detect_placeholders_in_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Detecta placeholders en texto plano y los convierte
        """
        placeholders = []
        # Buscar texto que podría ser placeholder (palabras en mayúsculas, entre corchetes, etc.)
        patterns = [
            r'\[([A-Z_]+)\]',  # [NOMBRE_PLACEHOLDER]
            r'\{([A-Z_]+)\}',  # {NOMBRE_PLACEHOLDER}
            r'([A-Z_]{3,})',   # PALABRAS_EN_MAYUSCULAS
        ]
        
        for pattern in patterns:
            matches = re.finditer(pattern, text)
            for match in matches:
                placeholder_name = match.group(1).lower()
                placeholders.append({
                    'name': placeholder_name,
                    'type': self._infer_placeholder_type(placeholder_name),
                    'position': match.span(),
                    'original_text': match.group(0)
                })
        
        return placeholders
    
    def _infer_placeholder_type(self, name: str) -> str:
        """
        Infiere el tipo de placeholder basado en su nombre
        """
        name_lower = name.lower()
        
        if any(word in name_lower for word in ['fecha', 'date', 'dia', 'mes', 'año']):
            return 'date'
        elif any(word in name_lower for word in ['numero', 'cantidad', 'precio', 'total', 'number']):
            return 'number'
        elif any(word in name_lower for word in ['email', 'correo', 'mail']):
            return 'email'
        elif any(word in name_lower for word in ['telefono', 'phone', 'tel', 'celular']):
            return 'phone'
        else:
            return 'text'
    
    # Métodos auxiliares para extraer estilos y propiedades
    
    def _get_alignment(self, alignment):
        """Convierte alineación de docx a CSS"""
        if alignment == WD_ALIGN_PARAGRAPH.CENTER:
            return 'center'
        elif alignment == WD_ALIGN_PARAGRAPH.RIGHT:
            return 'right'
        elif alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            return 'justify'
        else:
            return 'left'
    
    def _get_color(self, color):
        """Extrae color de texto"""
        if color and color.rgb:
            return f"#{color.rgb}"
        return None
    
    def _is_red_text(self, color):
        """Detecta si el texto es rojo (placeholder)"""
        if color and color.rgb:
            rgb_str = str(color.rgb)
            return rgb_str.startswith('FF0000') or rgb_str.startswith('ff0000')
        return False
    
    def _is_red_excel_cell(self, cell):
        """Detecta si una celda de Excel tiene texto rojo"""
        if cell.font and cell.font.color and cell.font.color.rgb:
            return str(cell.font.color.rgb).upper().startswith('FFFF0000')
        return False
    
    def _extract_excel_font(self, font):
        """Extrae información de fuente de Excel"""
        return {
            'name': font.name,
            'size': font.size,
            'bold': font.bold,
            'italic': font.italic,
            'color': str(font.color.rgb) if font.color and font.color.rgb else None
        }
    
    def _extract_excel_fill(self, fill):
        """Extrae información de relleno de Excel"""
        return {
            'pattern_type': fill.patternType,
            'fg_color': str(fill.fgColor.rgb) if fill.fgColor and fill.fgColor.rgb else None,
            'bg_color': str(fill.bgColor.rgb) if fill.bgColor and fill.bgColor.rgb else None
        }
    
    def _extract_excel_border(self, border):
        """Extrae información de bordes de Excel"""
        return {
            'left': border.left.style if border.left else None,
            'right': border.right.style if border.right else None,
            'top': border.top.style if border.top else None,
            'bottom': border.bottom.style if border.bottom else None
        }
    
    def _extract_excel_alignment(self, alignment):
        """Extrae información de alineación de Excel"""
        return {
            'horizontal': alignment.horizontal,
            'vertical': alignment.vertical,
            'wrap_text': alignment.wrap_text
        }
    
    def _paragraph_to_html(self, para_data: Dict) -> str:
        """Convierte datos de párrafo a HTML"""
        css_classes = ['paragraph']
        if para_data['alignment'] != 'left':
            css_classes.append(para_data['alignment'])
        
        html_content = []
        for run in para_data['runs']:
            run_html = self._run_to_html(run)
            html_content.append(run_html)
        
        return f'<p class="{" ".join(css_classes)}">{"".join(html_content)}</p>'
    
    def _run_to_html(self, run_data: Dict) -> str:
        """Convierte un run de texto a HTML"""
        text = run_data['text']
        css_classes = []
        
        if run_data['bold']:
            css_classes.append('bold')
        if run_data['italic']:
            css_classes.append('italic')
        if run_data['underline']:
            css_classes.append('underline')
        if run_data['is_placeholder']:
            css_classes.append('red-text')
            # Convertir texto rojo a placeholder
            text = self._convert_to_placeholder(text)
        
        if css_classes:
            return f'<span class="{" ".join(css_classes)}">{text}</span>'
        else:
            return text
    
    def _table_to_html(self, table_data: Dict) -> str:
        """Convierte datos de tabla a HTML"""
        html_parts = ['<table class="document-table">']
        
        for row in table_data['rows']:
            html_parts.append('<tr>')
            for cell in row['cells']:
                cell_content = []
                for para in cell['paragraphs']:
                    if para['text'].strip():
                        cell_content.append(self._paragraph_to_html(para))
                
                html_parts.append(f'<td>{"".join(cell_content) if cell_content else cell["text"]}</td>')
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return '\n'.join(html_parts)
    
    def _sheet_to_html_table(self, sheet_data: Dict) -> str:
        """Convierte datos de hoja Excel a tabla HTML"""
        if not sheet_data['cells']:
            return '<p>Hoja vacía</p>'
        
        # Determinar dimensiones
        max_row = sheet_data['dimensions']['max_row']
        max_col = sheet_data['dimensions']['max_column']
        
        html_parts = ['<table class="excel-table">']
        
        for row in range(1, max_row + 1):
            html_parts.append('<tr>')
            for col in range(1, max_col + 1):
                coord = f"{get_column_letter(col)}{row}"
                cell_data = sheet_data['cells'].get(coord, {'value': '', 'is_placeholder': False})
                
                cell_value = cell_data['value']
                css_classes = []
                
                if cell_data.get('is_placeholder'):
                    css_classes.append('placeholder')
                    cell_value = self._convert_to_placeholder(cell_value)
                
                class_attr = f' class="{" ".join(css_classes)}"' if css_classes else ''
                html_parts.append(f'<td{class_attr}>{cell_value}</td>')
            
            html_parts.append('</tr>')
        
        html_parts.append('</table>')
        return '\n'.join(html_parts)
    
    def _convert_to_placeholder(self, text: str) -> str:
        """Convierte texto a formato de placeholder"""
        # Limpiar y normalizar el texto
        clean_text = re.sub(r'[^\w\s]', '', text.strip())
        placeholder_name = clean_text.lower().replace(' ', '_')
        
        if placeholder_name:
            return f'<span class="placeholder" data-placeholder="{placeholder_name}">{{{{{placeholder_name}}}}}</span>'
        else:
            return text
    
    def _convert_text_to_html_with_placeholders(self, text: str) -> str:
        """Convierte texto plano a HTML detectando posibles placeholders"""
        paragraphs = text.split('\n\n')
        html_parts = []
        
        for para in paragraphs:
            if para.strip():
                # Detectar y convertir posibles placeholders
                para_html = self._detect_and_convert_placeholders_in_paragraph(para.strip())
                html_parts.append(f'<p>{para_html}</p>')
        
        return '\n'.join(html_parts)
    
    def _detect_and_convert_placeholders_in_paragraph(self, text: str) -> str:
        """Detecta y convierte placeholders en un párrafo"""
        # Patrones para detectar posibles placeholders
        patterns = [
            (r'\[([A-Z_]+)\]', r'<span class="placeholder" data-placeholder="\1">{{\1}}</span>'),
            (r'\{([A-Z_]+)\}', r'<span class="placeholder" data-placeholder="\1">{{\1}}</span>'),
            (r'([A-Z_]{3,})', r'<span class="placeholder" data-placeholder="\1">{{\1}}</span>')
        ]
        
        result = text
        for pattern, replacement in patterns:
            result = re.sub(pattern, replacement, result)
        
        return result

# Instancia global del servicio
document_parser = DocumentParserService()